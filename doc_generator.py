from docx import Document
import io


def generate_docx(summary_data, company_name, meeting_date):
    doc = Document()

    doc.add_heading(f"{company_name} Meeting Notes", level=0)
    doc.add_paragraph(f"Date: {meeting_date}", style="Heading 2").alignment = 2

    doc.add_heading("1. Minutes of the Meeting (MoM)", level=1)
    for line in summary_data["mom"]:
        doc.add_paragraph(line.strip(), style="List Bullet")

    doc.add_heading("2. To-Do List", level=1)
    for item in summary_data["todo_list"]:
        doc.add_paragraph(item.strip(), style="List Bullet")

    doc.add_heading("3. Action Points / Action Plan", level=1)
    section_titles = {
        "decision_made": "Key Decisions Made",
        "key_services_to_promote": "Key Services to Promote",
        "target_geography": "Target Geography",
        "budget_and_timeline": "Budget and Timeline",
        "lead_management_strategy": "Lead Management Strategy",
        "next_steps_and_ownership": "Next Steps and Ownership",
    }
    for key, title in section_titles.items():
        doc.add_heading(title, level=2)
        for item in summary_data["action_plan"].get(key, []):
            doc.add_paragraph(item.strip(), style="List Bullet")

    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)

    return docx_stream.read()
